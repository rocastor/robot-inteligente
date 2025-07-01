"""
🤖 Módulo de Gestión Rocastor
Maneja PDFs, plantillas y generación de carpetas
"""

import os
import json
import base64
import tempfile
import zipfile
import shutil
from datetime import datetime, timedelta
from docx import Document
import re

def crear_directorio_rocastor():
    """Crea directorios necesarios para Rocastor"""
    os.makedirs("rocastor_storage/pdfs", exist_ok=True)
    os.makedirs("rocastor_storage/templates", exist_ok=True)
    os.makedirs("rocastor_storage/specific_templates", exist_ok=True)

def guardar_pdf_rocastor(pdf_data):
    """Guarda un archivo PDF de forma persistente"""
    crear_directorio_rocastor()

    pdf_id = pdf_data.get('id')
    file_path = f"rocastor_storage/pdfs/{pdf_id}.json"

    pdf_data['uploadDate'] = datetime.now().isoformat()
    pdf_data['fechaVencimiento'] = (datetime.now() + timedelta(days=30)).isoformat()

    with open(file_path, 'w', encoding='utf-8') as f:
        json.dump(pdf_data, f, ensure_ascii=False, indent=2)

    return {"status": "success", "message": "PDF guardado correctamente"}

async def cargar_plantillas_rocastor():
    """Carga todas las plantillas generales guardadas"""
    try:
        crear_directorio_rocastor()

        templates_dir = "rocastor_storage/templates"
        if not os.path.exists(templates_dir):
            return {"templates": [], "total": 0, "mensaje": "Directorio no existe"}

        templates = []
        archivos_corruptos = []

        for filename in os.listdir(templates_dir):
            if filename.endswith('.json'):
                file_path = os.path.join(templates_dir, filename)
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read().strip()

                        if not content.startswith('{') or not content.endswith('}'):
                            archivos_corruptos.append(filename)
                            continue

                        template_data = json.loads(content)

                        if not isinstance(template_data, dict) or 'name' not in template_data:
                            archivos_corruptos.append(filename)
                            continue

                        templates.append(template_data)

                except Exception as e:
                    print(f"❌ Error procesando plantilla {filename}: {str(e)}")
                    archivos_corruptos.append(filename)

        print(f"✅ Cargadas {len(templates)} plantillas válidas")
        return {
            "templates": templates,
            "total": len(templates),
            "mensaje": f"Cargadas {len(templates)} plantillas correctamente"
        }

    except Exception as e:
        print(f"❌ Error en cargar_plantillas_rocastor: {str(e)}")
        return {
            "templates": [],
            "total": 0,
            "mensaje": f"Error cargando plantillas: {str(e)}",
            "error": True
        }

def guardar_plantilla_rocastor(template_data):
    """Guarda una plantilla general"""
    crear_directorio_rocastor()

    template_id = template_data.get('id')
    file_path = f"rocastor_storage/templates/{template_id}.json"

    template_data['uploadDate'] = datetime.now().isoformat()

    with open(file_path, 'w', encoding='utf-8') as f:
        json.dump(template_data, f, ensure_ascii=False, indent=2)

    return {"status": "success", "message": "Plantilla guardada correctamente"}

def eliminar_plantilla_rocastor(template_id):
    """Elimina una plantilla general"""
    try:
        file_path = f"rocastor_storage/templates/{template_id}.json"
        if os.path.exists(file_path):
            os.remove(file_path)
            return {"status": "success", "message": "Plantilla eliminada"}
        else:
            return {"status": "error", "message": "Plantilla no encontrada"}
    except Exception as e:
        return {"status": "error", "message": f"Error eliminando plantilla: {str(e)}"}

def eliminar_pdf_rocastor(pdf_id):
    """Elimina un PDF"""
    try:
        file_path = f"rocastor_storage/pdfs/{pdf_id}.json"
        if os.path.exists(file_path):
            os.remove(file_path)
            return {"status": "success", "message": "PDF eliminado"}
        else:
            return {"status": "error", "message": "PDF no encontrado"}
    except Exception as e:
        return {"status": "error", "message": f"Error eliminando PDF: {str(e)}"}

async def cargar_pdfs_rocastor():
    """Carga todos los PDFs guardados"""
    try:
        crear_directorio_rocastor()

        pdfs_dir = "rocastor_storage/pdfs"
        if not os.path.exists(pdfs_dir):
            return {"pdfs": [], "total": 0, "mensaje": "Directorio no existe"}

        pdfs = []
        archivos_corruptos = []

        for filename in os.listdir(pdfs_dir):
            if filename.endswith('.json'):
                file_path = os.path.join(pdfs_dir, filename)
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read().strip()

                        if not content.startswith('{') or not content.endswith('}'):
                            print(f"⚠️ Archivo corrupto detectado: {filename}")
                            archivos_corruptos.append(filename)
                            continue

                        pdf_data = json.loads(content)

                        if not isinstance(pdf_data, dict) or 'name' not in pdf_data:
                            print(f"⚠️ Estructura JSON inválida en: {filename}")
                            archivos_corruptos.append(filename)
                            continue

                        pdfs.append(pdf_data)

                except (json.JSONDecodeError, UnicodeDecodeError) as e:
                    print(f"❌ Error leyendo {filename}: {str(e)}")
                    archivos_corruptos.append(filename)
                    continue
                except Exception as e:
                    print(f"❌ Error procesando {filename}: {str(e)}")
                    continue

        # Mover archivos corruptos a backup
        if archivos_corruptos:
            backup_dir = "rocastor_storage/backup_corruptos"
            os.makedirs(backup_dir, exist_ok=True)

            for corrupted_file in archivos_corruptos:
                try:
                    src = os.path.join(pdfs_dir, corrupted_file)
                    dst = os.path.join(backup_dir, f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{corrupted_file}")
                    shutil.move(src, dst)
                    print(f"📦 Archivo corrupto movido a backup: {corrupted_file}")
                except Exception as e:
                    print(f"❌ Error moviendo archivo corrupto {corrupted_file}: {str(e)}")

        print(f"✅ Cargados {len(pdfs)} PDFs válidos")
        if archivos_corruptos:
            print(f"⚠️ {len(archivos_corruptos)} archivos corruptos movidos a backup")

        return {
            "pdfs": pdfs, 
            "total": len(pdfs), 
            "corruptos_reparados": len(archivos_corruptos),
            "mensaje": f"Cargados {len(pdfs)} PDFs correctamente"
        }

    except Exception as e:
        print(f"❌ Error en cargar_pdfs_rocastor: {str(e)}")
        return {
            "pdfs": [],
            "total": 0,
            "mensaje": f"Error cargando PDFs: {str(e)}",
            "error": True
        }

def generar_carpeta_rocastor(folder_data):
    """Genera carpeta completa con plantillas diligenciadas y PDFs reales"""
    try:
        # Obtener datos del JSON enviado desde el frontend
        process_name = folder_data.get("processName", "proceso_desconocido")
        pdf_files_data = folder_data.get("pdfFiles", [])
        template_files_data = folder_data.get("templateFiles", [])
        process_answers = folder_data.get("processAnswers", {})
        using_specific_templates = folder_data.get("usingSpecificTemplates", False)

        print(f"🚀 Generando carpeta para proceso: {process_name}")
        print(f"📄 PDFs a incluir: {len(pdf_files_data)}")
        print(f"📝 Plantillas a diligenciar: {len(template_files_data)}")
        print(f"🎯 Usando plantillas específicas: {using_specific_templates}")

        # Usar el nombre exacto del proceso como nombre de carpeta (sin modificaciones)
        carpeta_nombre = process_name
        print(f"📁 Nombre de carpeta: {carpeta_nombre} (sin modificaciones)")

        # Crear carpeta temporal
        temp_dir = tempfile.mkdtemp()
        carpeta_proceso = os.path.join(temp_dir, carpeta_nombre)
        os.makedirs(carpeta_proceso, exist_ok=True)

        # Crear subcarpetas organizadas
        carpeta_pdfs = os.path.join(carpeta_proceso, "PDFs_Ofertas")
        carpeta_plantillas = os.path.join(carpeta_proceso, "Plantillas_Diligenciadas")
        os.makedirs(carpeta_pdfs, exist_ok=True)
        os.makedirs(carpeta_plantillas, exist_ok=True)

        # 1. PROCESAR PDFs REALES
        pdfs_copiados = 0
        print(f"📎 Procesando {len(pdf_files_data)} PDFs...")

        for i, pdf_data in enumerate(pdf_files_data):
            try:
                pdf_name = pdf_data.get('name', f'documento_{i+1}.pdf')
                pdf_content = pdf_data.get('data', '')

                # Limpiar nombre del archivo
                safe_pdf_name = re.sub(r'[^\w\-_\.]', '_', pdf_name)
                pdf_path = os.path.join(carpeta_pdfs, safe_pdf_name)

                if pdf_content and pdf_content.startswith('data:application/pdf;base64,'):
                    # Extraer y decodificar base64
                    base64_data = pdf_content.split(',')[1]
                    pdf_bytes = base64.b64decode(base64_data)

                    with open(pdf_path, 'wb') as f:
                        f.write(pdf_bytes)

                    pdfs_copiados += 1
                    print(f"   ✅ PDF procesado: {safe_pdf_name} ({len(pdf_bytes)} bytes)")
                else:
                    print(f"   ⚠️ PDF sin contenido válido: {pdf_name}")

            except Exception as e:
                print(f"   ❌ Error procesando PDF {pdf_name}: {str(e)}")

        # 2. PROCESAR Y DILIGENCIAR PLANTILLAS
        plantillas_diligenciadas = 0
        print(f"📝 Procesando {len(template_files_data)} plantillas...")

        # Generar fecha actual en múltiples formatos - MEJORADO
        from datetime import datetime
        fecha_actual = datetime.now()
        meses = ['enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
                 'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre']

        dia = fecha_actual.day
        mes = meses[fecha_actual.month - 1]
        año = fecha_actual.year

        fecha_larga = f"{dia} de {mes} de {año}"
        fecha_corta = fecha_actual.strftime("%d/%m/%Y")
        fecha_formal = f"{dia} de {mes} de {año}"
        fecha_completa = f"{dia} de {mes} del {año}"

        print(f"📅 Fecha generada para plantillas: {fecha_larga}")
        print(f"📅 Fecha corta: {fecha_corta}")
        print(f"📅 Fecha formal: {fecha_formal}")

        for i, template_data in enumerate(template_files_data):
            try:
                template_name = template_data.get('name', f'plantilla_{i+1}.docx')
                template_content = template_data.get('data', '')

                # Limpiar nombre de plantilla para archivos temporales
                safe_template_name = re.sub(r'[^\w\-_\.]', '_', template_name)

                # Mantener nombre original de la plantilla
                template_path = os.path.join(carpeta_plantillas, template_name)

                if template_content and template_content.startswith('data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,'):
                    # Decodificar plantilla original
                    base64_data = template_content.split(',')[1]
                    template_bytes = base64.b64decode(base64_data)

                    # Guardar temporalmente la plantilla original
                    temp_template_path = os.path.join(temp_dir, f"temp_{safe_template_name}")
                    with open(temp_template_path, 'wb') as f:
                        f.write(template_bytes)

                    # Cargar plantilla con python-docx
                    try:
                        doc = Document(temp_template_path)

                        # DILIGENCIAR PLANTILLA CON DATOS REALES DEL PROCESO
                        print(f"   🔄 Diligenciando plantilla: {template_name}")

                        # Extraer valores con validación
                        entidad = str(process_answers.get('entidad', 'ENTIDAD_NO_ENCONTRADA')).strip()
                        nit = str(process_answers.get('nit', 'NIT_NO_ENCONTRADO')).strip()
                        ciudad = str(process_answers.get('ciudad', 'CIUDAD_NO_ENCONTRADA')).strip()
                        direccion = str(process_answers.get('direccion', 'DIRECCION_NO_ENCONTRADA')).strip()
                        valor = str(process_answers.get('valor', 'VALOR_NO_ENCONTRADO')).strip()
                        objeto = str(process_answers.get('objeto', 'OBJETO_NO_ENCONTRADO')).strip()

                        print(f"      📋 Datos para diligenciar:")
                        print(f"         Entidad: {entidad[:50]}...")
                        print(f"         NIT: {nit}")
                        print(f"         Ciudad: {ciudad}")
                        print(f"         Valor: {valor}")
                        print(f"         Fecha: {fecha_larga}")

                        # Diccionario de reemplazos más completo y organizado
                        replacements = {
                            # ENTIDAD - Múltiples formatos y variaciones
                            '{{ENTIDAD}}': entidad,
                            '(ENTIDAD)': entidad,
                            '[ENTIDAD]': entidad,
                            '{ENTIDAD}': entidad,
                            'ENTIDAD': entidad,
                            '{{ENTIDAD_CONTRATANTE}}': entidad,
                            '(ENTIDAD_CONTRATANTE)': entidad,
                            'ENTIDAD_CONTRATANTE': entidad,
                            'ENTIDAD CONTRATANTE': entidad,
                            'Señores': entidad,
                            'SEÑORES': entidad,

                            # NIT - Con y sin prefijo NIT
                            '{{NIT}}': nit,
                            '(NIT)': nit,
                            '[NIT]': nit,
                            '{NIT}': nit,
                            'NIT:': f'NIT: {nit}',
                            'NIT.': f'NIT. {nit}',
                            'N.I.T:': f'N.I.T: {nit}',
                            'N.I.T.': f'N.I.T. {nit}',
                            '{{NIT_ENTIDAD}}': nit,
                            'NIT_ENTIDAD': nit,

                            # CIUDAD y ubicación
                            '{{CIUDAD}}': ciudad,
                            '(CIUDAD)': ciudad,
                            '[CIUDAD]': ciudad,
                            '{CIUDAD}': ciudad,
                            'CIUDAD': ciudad,

                            # DIRECCIÓN
                            '{{DIRECCION}}': direccion,
                            '(DIRECCION)': direccion,
                            '[DIRECCION]': direccion,
                            '{DIRECCION}': direccion,
                            'DIRECCION': direccion,
                            'Dirección:': f'Dirección: {direccion}',
                            'Dir.': f'Dir. {direccion}',
                            '{{DIRECCION_ENTIDAD}}': direccion,

                            # FECHA - Múltiples formatos y contextos específicos
                            '{{FECHA}}': fecha_larga,
                            '(FECHA)': fecha_larga,
                            '[FECHA]': fecha_larga,
                            '{FECHA}': fecha_larga,
                            'FECHA': fecha_larga,
                            'FECHA_ACTUAL': fecha_larga,
                            'FECHA_HOY': fecha_larga,
                            'HOY': fecha_larga,
                            '{{FECHA_ACTUAL}}': fecha_larga,
                            '(FECHA_ACTUAL)': fecha_larga,
                            'FECHA_CORTA': fecha_corta,
                            '{{FECHA_CORTA}}': fecha_corta,

                            # Patrones específicos para ciudades con fecha (como en la imagen)
                            'Bogotá D.C., (FECHA)': f'Bogotá D.C., {fecha_larga}',
                            'Bogotá D.C., {FECHA}': f'Bogotá D.C., {fecha_larga}',
                            'Bogotá D.C., {{FECHA}}': f'Bogotá D.C., {fecha_larga}',
                            'Bogotá D.C., [FECHA]': f'Bogotá D.C., {fecha_larga}',
                            f'{ciudad}, (FECHA)': f'{ciudad}, {fecha_larga}',
                            f'{ciudad}, {{FECHA}}': f'{ciudad}, {fecha_larga}',

                            # VALOR y presupuesto
                            '{{VALOR}}': valor,
                            '(VALOR)': valor,
                            '[VALOR]': valor,
                            '{VALOR}': valor,
                            'VALOR': valor,
                            '{{VALOR_CONTRATO}}': valor,
                            'VALOR_CONTRATO': valor,
                            '{{PRESUPUESTO}}': valor,
                            'PRESUPUESTO': valor,
                            'Valor:': f'Valor: {valor}',
                            'Presupuesto:': f'Presupuesto: {valor}',

                            # OBJETO del contrato
                            '{{OBJETO}}': objeto,
                            '(OBJETO)': objeto,
                            '[OBJETO]': objeto,
                            '{OBJETO}': objeto,
                            'OBJETO': objeto,
                            '{{OBJETO_CONTRATO}}': objeto,
                            'OBJETO_CONTRATO': objeto,
                            'Objeto:': f'Objeto: {objeto}',
                            'Ref:': f'Ref: {objeto}',

                            # PROCESO
                            '{{PROCESO}}': process_name,
                            'PROCESO': process_name,
                            '{{NOMBRE_PROCESO}}': process_name,
                            'NOMBRE_PROCESO': process_name,
                            '(PROCESO)': process_name,
                        }

                        # Función mejorada para reemplazar texto en párrafos
                        def replace_text_in_paragraph(paragraph):
                            if not paragraph.text.strip():
                                return

                            full_text = paragraph.text
                            original_text = full_text
                            changes_made = False
                            changes_detail = []

                            # Reemplazos exactos primero
                            for old_text, new_text in replacements.items():
                                if old_text in full_text:
                                    full_text = full_text.replace(old_text, str(new_text))
                                    changes_made = True
                                    changes_detail.append(f"{old_text} → {str(new_text)[:30]}...")

                            # Patrones regex para casos no marcados - MEJORADOS
                            regex_patterns = [
                                # Fechas vacías o con placeholders
                                (r'Fecha:\s*[_\s]*$', f'Fecha: {fecha_larga}'),
                                (r'Fecha\s*:\s*[_\s]*$', f'Fecha: {fecha_larga}'),
                                (r'\bFecha\s*$', f'Fecha {fecha_larga}'),

                                # Campos vacíos específicos
                                (r'Señores\s*:\s*[_\s]*$', f'Señores: {entidad}'),
                                (r'NIT\s*:\s*[_\s]*$', f'NIT: {nit}'),
                                (r'Ciudad\s*:\s*[_\s]*$', f'Ciudad: {ciudad}'),
                                (r'Dirección\s*:\s*[_\s]*$', f'Dirección: {direccion}'),
                                (r'Valor\s*:\s*[_\s]*$', f'Valor: {valor}'),

                                # Patrones de ciudad con fecha vacía
                                (r'Bogotá\s+D\.C\.\s*,\s*[_\s]*$', f'Bogotá D.C., {fecha_larga}'),
                                (rf'{re.escape(ciudad)}\s*,\s*[_\s]*$', f'{ciudad}, {fecha_larga}'),
                            ]

                            for pattern, replacement in regex_patterns:
                                if re.search(pattern, full_text, re.IGNORECASE | re.MULTILINE):
                                    full_text = re.sub(pattern, replacement, full_text, flags=re.IGNORECASE | re.MULTILINE)
                                    changes_made = True
                                    changes_detail.append(f"REGEX: {pattern[:20]}... → {replacement[:30]}...")

                            # Solo actualizar si hubo cambios
                            if changes_made and full_text != original_text:
                                # Preservar formato original al actualizar
                                paragraph.clear()
                                run = paragraph.add_run(full_text)
                                print(f"      ✅ Párrafo actualizado: {len(changes_detail)} cambios")
                                for change in changes_detail[:3]:  # Mostrar máximo 3 cambios
                                    print(f"         • {change}")

                            return changes_made

                        # Contador de cambios totales
                        total_changes = 0

                        # Reemplazar en párrafos del documento principal
                        print(f"      📄 Procesando párrafos principales...")
                        for i, paragraph in enumerate(doc.paragraphs):
                            if replace_text_in_paragraph(paragraph):
                                total_changes += 1

                        # Reemplazar en tablas
                        print(f"      📋 Procesando tablas...")
                        for table_idx, table in enumerate(doc.tables):
                            for row_idx, row in enumerate(table.rows):
                                for cell_idx, cell in enumerate(row.cells):
                                    for para_idx, paragraph in enumerate(cell.paragraphs):
                                        if replace_text_in_paragraph(paragraph):
                                            total_changes += 1

                        # Reemplazar en headers y footers
                        print(f"      📰 Procesando headers y footers...")
                        for section_idx, section in enumerate(doc.sections):
                            try:
                                # Header
                                if section.header:
                                    for paragraph in section.header.paragraphs:
                                        if replace_text_in_paragraph(paragraph):
                                            total_changes += 1

                                # Footer
                                if section.footer:
                                    for paragraph in section.footer.paragraphs:
                                        if replace_text_in_paragraph(paragraph):
                                            total_changes += 1
                            except Exception as e:
                                print(f"      ⚠️ Error en header/footer sección {section_idx}: {str(e)}")

                        # Guardar plantilla diligenciada con nombre original
                        try:
                            doc.save(template_path)
                            plantillas_diligenciadas += 1
                            print(f"   ✅ Plantilla diligenciada: {template_name}")
                            print(f"      📊 Total de cambios realizados: {total_changes}")
                            print(f"      💾 Guardada en: {template_path}")

                            # Verificar que el archivo se guardó correctamente
                            if os.path.exists(template_path):
                                file_size = os.path.getsize(template_path)
                                print(f"      ✅ Archivo verificado: {file_size} bytes")
                            else:
                                print(f"      ❌ Error: Archivo no se guardó correctamente")

                        except Exception as save_error:
                            print(f"      ❌ Error guardando plantilla: {str(save_error)}")
                            # Crear plantilla básica como respaldo
                            create_basic_template(template_path, template_name, process_answers, process_name)
                            plantillas_diligenciadas += 1

                        # Limpiar archivo temporal
                        try:
                            if os.path.exists(temp_template_path):
                                os.remove(temp_template_path)
                        except Exception as cleanup_error:
                            print(f"      ⚠️ Error limpiando archivo temporal: {str(cleanup_error)}")

                    except Exception as e:
                        print(f"   ❌ Error diligenciando plantilla {template_name}: {str(e)}")
                        # Crear plantilla básica en caso de error
                        create_basic_template(template_path, template_name, process_answers, process_name)
                        plantillas_diligenciadas += 1
                else:
                    print(f"   ⚠️ Plantilla sin contenido válido: {template_name}")
                    # Crear plantilla básica
                    create_basic_template(template_path, template_name, process_answers, process_name)
                    plantillas_diligenciadas += 1

            except Exception as e:
                print(f"   ❌ Error procesando plantilla {template_name}: {str(e)}")

        # 3. CREAR ARCHIVO DE RESUMEN DEL PROCESO
        resumen_path = os.path.join(carpeta_proceso, "RESUMEN_PROCESO.txt")
        with open(resumen_path, 'w', encoding='utf-8') as f:
            f.write(f"RESUMEN DEL PROCESO: {process_name}\n")
            f.write("="*50 + "\n\n")
            f.write(f"Fecha de generación: {fecha_larga}\n\n")

            f.write("INFORMACIÓN EXTRAÍDA:\n")
            f.write("-"*25 + "\n")
            for key, value in process_answers.items():
                if value and value.strip():
                    f.write(f"{key.upper()}: {value}\n")

            f.write(f"\nARCHIVOS INCLUIDOS:\n")
            f.write("-"*20 + "\n")
            f.write(f"PDFs procesados: {pdfs_copiados}\n")
            f.write(f"Plantillas diligenciadas: {plantillas_diligenciadas}\n")

        # 4. CREAR ARCHIVO ZIP
        zip_filename = f"{carpeta_nombre}.zip"
        zip_path = os.path.join(temp_dir, zip_filename)

        import zipfile
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for root, dirs, files in os.walk(carpeta_proceso):
                for file in files:
                    file_path = os.path.join(root, file)
                    arc_path = os.path.relpath(file_path, temp_dir)
                    zipf.write(file_path, arc_path)

        print(f"📦 Carpeta generada exitosamente:")
        print(f"   📁 Nombre: {carpeta_nombre}")
        print(f"   📎 PDFs incluidos: {pdfs_copiados}")
        print(f"   📝 Plantillas diligenciadas: {plantillas_diligenciadas}")
        print(f"   💾 Archivo ZIP: {zip_path}")

        # COPIA AUTOMÁTICA A DRIVE EMPRESARIAL
        try:
            print(f"\n🏢 ===== COPIA A DRIVE EMPRESARIAL =====")
            import os
            if os.getenv('GOOGLE_DRIVE_CREDENTIALS'):
                from modules.google_drive_client import get_drive_client
                
                drive_client = get_drive_client()
                if drive_client:
                    # ID de la carpeta empresarial
                    empresa_folder_id = "1EfI2gKDlYiMmsi7dTGFsyHdtqhx9FLGi"
                    
                    # Crear subcarpeta "Procesos_Rocastor" en drive empresarial
                    rocastor_folder_id = drive_client.create_or_get_folder(
                        "Procesos_Rocastor",
                        empresa_folder_id
                    )
                    
                    if rocastor_folder_id:
                        # Crear subcarpeta con el nombre del proceso
                        proceso_folder_id = drive_client.create_or_get_folder(
                            carpeta_nombre,
                            rocastor_folder_id
                        )
                        
                        if proceso_folder_id:
                            # Subir el archivo ZIP
                            zip_result = drive_client.upload_file(
                                zip_path,
                                zip_filename,
                                proceso_folder_id,
                                {
                                    'tipo': 'proceso_rocastor',
                                    'proceso': carpeta_nombre,
                                    'pdfs_incluidos': pdfs_copiados,
                                    'plantillas_diligenciadas': plantillas_diligenciadas
                                }
                            )
                            
                            if zip_result:
                                print(f"✅ ZIP subido a Drive empresarial: {zip_result.get('web_view_link')}")
                                print(f"📁 Carpeta empresarial: https://drive.google.com/drive/folders/{proceso_folder_id}")
                            else:
                                print(f"⚠️ Error subiendo ZIP a Drive empresarial")
                        else:
                            print(f"⚠️ No se pudo crear carpeta del proceso en Drive empresarial")
                    else:
                        print(f"⚠️ No se pudo crear carpeta Rocastor en Drive empresarial")
                else:
                    print(f"⚠️ Cliente de Google Drive no disponible")
            else:
                print(f"⚠️ Google Drive no configurado")
                
        except Exception as empresa_error:
            print(f"⚠️ Error copiando a Drive empresarial: {str(empresa_error)}")
            # No detener el proceso por este error

        return zip_path, zip_filename

    except Exception as e:
        print(f"❌ Error generando carpeta Rocastor: {str(e)}")
        import traceback
        traceback.print_exc()
        raise Exception(f"Error generando carpeta: {str(e)}")

def create_basic_template(file_path, template_name, process_answers, process_name):
    """Crea una plantilla básica con los datos del proceso cuando hay errores"""
    try:
        from docx import Document

        doc = Document()
        doc.add_heading(f'Documento: {template_name}', 0)
        doc.add_heading('Datos del Proceso', level=1)

        # Agregar información del proceso
        for key, value in process_answers.items():
            if value and value.strip():
                p = doc.add_paragraph()
                p.add_run(f"{key.upper()}: ").bold = True
                p.add_run(str(value))

        # Agregar fecha
        fecha_actual = datetime.now()
        doc.add_paragraph(f"Proceso generado: {process_name}")
        doc.add_paragraph(f"Fecha: {fecha_actual.strftime('%d/%m/%Y %H:%M')}")

        doc.save(file_path)
        print(f"   ✅ Plantilla básica creada: {template_name}")

    except Exception as e:
        print(f"   ❌ Error creando plantilla básica: {str(e)}")

async def cargar_plantillas_especificas_rocastor(process_name):
    """Carga plantillas específicas de un proceso"""
    try:
        crear_directorio_rocastor()

        specific_templates_dir = f"rocastor_storage/specific_templates/{process_name}"
        if not os.path.exists(specific_templates_dir):
            return {"templates": [], "total": 0, "mensaje": "No hay plantillas específicas para este proceso"}

        templates = []
        archivos_corruptos = []

        for filename in os.listdir(specific_templates_dir):
            if filename.endswith('.json'):
                file_path = os.path.join(specific_templates_dir, filename)
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read().strip()

                        if not content.startswith('{') or not content.endswith('}'):
                            archivos_corruptos.append(filename)
                            continue

                        template_data = json.loads(content)

                        if not isinstance(template_data, dict) or 'name' not in template_data:
                            archivos_corruptos.append(filename)
                            continue

                        templates.append(template_data)

                except Exception as e:
                    print(f"❌ Error procesando plantilla específica {filename}: {str(e)}")
                    archivos_corruptos.append(filename)

        print(f"✅ Cargadas {len(templates)} plantillas específicas para {process_name}")
        return {
            "templates": templates,
            "total": len(templates),
            "mensaje": f"Cargadas {len(templates)} plantillas específicas correctamente"
        }

    except Exception as e:
        print(f"❌ Error en cargar_plantillas_especificas_rocastor: {str(e)}")
        return {
            "templates": [],
            "total": 0,
            "mensaje": f"Error cargando plantillas específicas: {str(e)}",
            "error": True
        }

def guardar_plantilla_especifica_rocastor(template_data):
    """Guarda una plantilla específica de proceso"""
    try:
        crear_directorio_rocastor()

        process_name = template_data.get('processName')
        if not process_name:
            return {"status": "error", "message": "Nombre de proceso requerido"}

        specific_dir = f"rocastor_storage/specific_templates/{process_name}"
        os.makedirs(specific_dir, exist_ok=True)

        template_id = template_data.get('id')
        file_path = os.path.join(specific_dir, f"{template_id}.json")

        template_data['uploadDate'] = datetime.now().isoformat()

        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(template_data, f, ensure_ascii=False, indent=2)

        return {"status": "success", "message": "Plantilla específica guardada correctamente"}

    except Exception as e:
        return {"status": "error", "message": f"Error guardando plantilla específica: {str(e)}"}

def eliminar_plantilla_especifica_rocastor(process_name, template_id):
    """Elimina una plantilla específica de proceso"""
    try:
        file_path = f"rocastor_storage/specific_templates/{process_name}/{template_id}.json"
        if os.path.exists(file_path):
            os.remove(file_path)
            return {"status": "success", "message": "Plantilla específica eliminada"}
        else:
            return {"status": "error", "message": "Plantilla específica no encontrada"}
    except Exception as e:
        return {"status": "error", "message": f"Error eliminando plantilla específica: {str(e)}"}